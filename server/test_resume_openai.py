import os
import re
import docx
import PyPDF2
from dotenv import load_dotenv
from collections import Counter
from tkinter import Tk
from tkinter.filedialog import askopenfilename

try:
    from openai import OpenAI
    openai_available = True
except ImportError:
    openai_available = False

# === Load API Key ===
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")

# === Text Extraction ===

def extract_text_from_pdf(file_path):
    text = ""
    try:
        with open(file_path, "rb") as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"[PDF ERROR] {e}")
    return text

def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
    except Exception as e:
        print(f"[DOCX ERROR] {e}")
    return text

def extract_text(file_path):
    if file_path.lower().endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.lower().endswith('.docx'):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format. Use PDF or DOCX.")

# === Keyword Extraction ===

def clean_and_tokenize(text):
    text = re.sub(r'[^a-zA-Z\s]', ' ', text)
    words = text.lower().split()
    stop_words = set([
        "a", "an", "the", "and", "or", "for", "in", "on", "at", "to", "of", "with", "is", "are", "was", "were", "i", "you",
        "he", "she", "it", "we", "they", "them", "this", "that", "these", "those", "have", "has", "had", "as", "by", "be",
        "will", "would", "can", "could", "should", "from", "my", "your", "their", "but", "not", "if", "than", "so", "then",
        "all", "any", "no", "just", "do", "does", "did", "also", "each", "more", "most", "such"
    ])
    return [word for word in words if word not in stop_words and len(word) > 2]

def extract_keywords(text, top_n=25):
    tokens = clean_and_tokenize(text)
    counter = Counter(tokens)
    return counter.most_common(top_n)

# === OpenAI Analysis ===

def analyze_with_openai(resume_text, job_text=None):
    if not openai_available or not api_key:
        return "⚠️ OpenAI analysis skipped (no key or library)."

    try:
        client = OpenAI(api_key=api_key)

        if job_text:
            prompt = (
                "You are a resume screening assistant.\n\n"
                f"Here is the job description:\n{job_text}\n\n"
                f"And here is the resume:\n{resume_text}\n\n"
                "Compare the resume to the job description. \n" 
                "Rate out of (0-100%) how well the past experience matches the job description which includes education and past job roles. Take other niche things like certifications and bootcamps into account.\n"
                "Rate out of (0-100%) how well the skills match the job description as well as highlight all already obtained relevant skills as well as a;; missing skills needed.\n" 
                "Rate out of (0-100%) how well the Structure of the resume is and point out improvements whether it be spelling/grammer, repition, formatting, readability, keyword useage, etc.\n"
                ""
            )
        else:
            prompt = f"Analyze this resume:\n\n{resume_text}"

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt}
            ],
        )
        return response.choices[0].message.content

    except Exception as e:
        return f"❌ OpenAI Error: {e}"

# === File Picker ===

def pick_file():
    root = Tk()
    root.withdraw()
    return askopenfilename(
        title="Select Resume File",
        filetypes=[("PDF files", "*.pdf"), ("Word Documents", "*.docx")]
    )

# === Main ===

def main():
    print("📁 Select a resume file (PDF or DOCX)...")
    file_path = pick_file()

    if not file_path:
        print("❌ No file selected.")
        return

    print(f"📄 Extracting text from {os.path.basename(file_path)}...")
    resume_text = extract_text(file_path)

    if not resume_text.strip():
        print("⚠️ No text found in file.")
        return

    print("📊 Extracting keywords (local frequency)...")
    keywords = extract_keywords(resume_text)

    print("\n🔑 Top Keywords (Local):\n")
    print(f"{'KEYWORD':<20} {'FREQUENCY'}")
    print("-" * 30)
    for word, freq in keywords:
        print(f"{word:<20} {freq}")

    # Ask for job description paste
    print("\n📝 Do you want to paste a job description to compare with this resume? (y/n)")
    if input("> ").strip().lower() == "y":
        print("\n📋 Paste job description (end with empty line):")
        lines = []
        while True:
            line = input()
            if not line.strip():
                break
            lines.append(line)
        job_description = "\n".join(lines)
    else:
        job_description = None

    # Run OpenAI analysis
    if openai_available and api_key:
        print("\n🧠 Sending resume", end='')
        if job_description:
            print(" + job description", end='')
        print(" to OpenAI for analysis...")
        result = analyze_with_openai(resume_text, job_description)
        print("\n📬 OpenAI Response:\n")
        print(result)
    else:
        print("\n💡 Skipping OpenAI analysis (missing API key or library).")

if __name__ == "__main__":
    main()
