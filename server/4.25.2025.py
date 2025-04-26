from flask import Flask, jsonify, request
from flask_cors import CORS
import psycopg2
from dotenv import load_dotenv, dotenv_values
from openai import OpenAI
import json
import os
import PyPDF2
import docx

app = Flask(__name__)
cors = CORS(app, origin='*')

load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
psql_password = os.getenv("PSQL_PASSWORD")

def get_db_connection():
    return psycopg2.connect(
        host="localhost",
        dbname="postgres",
        user="postgres",
        password=psql_password,
        port=5432
    )

def openai(context, prompt):

    client = OpenAI(api_key=api_key)

    completion = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": context
            },
            {
                "role": "user",
                "content": prompt,
            },
        ],
    )
    response = completion.choices[0].message.content
    clean_response = response.strip()
    clean_response = response.strip("```json")
    return clean_response

def extract_text_from_pdf(file_stream):
    text = ""
    try:
        pdf_reader = PyPDF2.PdfReader(file_stream)
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        print(f"[PDF ERROR] {e}")
    return text

def extract_text_from_docx(file_stream):
    text = ""
    try:
        doc = docx.Document(file_stream)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
    except Exception as e:
        print(f"[DOCX ERROR] {e}")
    return text

def extract_text(file_stream, filename):
    if filename.lower().endswith('.pdf'):
        return extract_text_from_pdf(file_stream)
    elif filename.lower().endswith('.docx'):
        return extract_text_from_docx(file_stream)
    else:
        raise ValueError("Unsupported file format. Use PDF or DOCX.")
    
@app.route("/api/upload-resume", methods = ['POST'])
def upload_resume():
    resume = request.files["Resume"]

    return resume.filename


@app.route("/api/upload-job-description", methods = ['POST'])
def upload_job_description():
    data = request.get_json()
    job_description = data.get("job_description")

    return job_description

@app.route("/api/resume-score", methods = ['GET'])
def calculate_score():
    data = request.get_json()
    session_id = data.get("session_id")

    connection = get_db_connection()
    cursor = connection.cursor()

    cursor.execute("SELECT session_id, resume_scores FROM scores WHERE session_id = %s", (session_id))
    scores = cursor.fetchone()
    if (scores):
        return jsonify(scores)
    else:
        context = f"""
        You are an AI expert resume analyzer that always responds in clean nested JSON format. The user will input a resume, in text form, 
        and a job description, in  text form. You must analyze the resume based on the job descriptions.
        Your analyses will return the following.
        experience_score (0-100): Based on work experience and education keywords
        skills_score (0-100): Based on soft skills and technical skills keywords
        structure_score (0-100): Based on spelling & grammer, repetition, format, readability and keyword usage
        For each score alos include a list of positives and negatives.

        Example Output:
        [
            {{
                "experience_score": 85,
                "experience_positives": ["work experience"],
                "experience_negatives": ["education"],

                "skills_score": 100,
                "skills_positives": ["soft skills", "technical skills"],
                "skills_negatives": [],

                "structure_score": 75,
                "structure_positives": ["spelling & grammer", "repetition", "format"],
                "structure_negatives": ["readability", "keyword usage"]
            }},
            ...
        ]

        General Guidelines:
        1. Use ONLY the following exact labels for positives/negatives:
            - Experience: "work experience", "education"
            - Skills: "soft skills", "technical skills"
            - Structure: "spelling & grammar", "repetition", "format", "readability", "keyword usage"
        2. Do not paraphrase or summarize - only use the exact category labels
        3. Only respond with the single JSON object, not an array and no additional comments, etc.
        """

        cursor.execute("SELECT job_description FROM job_description WHERE session_id = %s", (session_id,))
        job_description = cursor.fetchone()[0]

        cursor.execute("SELECT resume_text FROM resumes WHERE session_id = %s", (session_id,))
        resume = cursor.fetchone()[0]

        prompt = f"""
        -RESUME-
        {resume}

        -JOB DESCIPTION-
        {job_description}
        """

        response = openai(context, prompt)
        json_response = json.loads(response)

        cursor.execute(
        "INSERT INTO scores (session_id, resume_scores) VALUES (%s, %s)",
        (session_id, json.dumps(json_response)))
     
        connection.commit()
        cursor.close()
        connection.close()

        return jsonify(json_response)

@app.route("/api/resume-improvements", methods = ['GET'])
def resume_improvements():
    session_id = 1 #hard coded for now

    connection = get_db_connection()
    cursor = connection.cursor()

    cursor.execute("SELECT improvements FROM resume_improvements WHERE session_id = %s", (session_id,))
    improvements = cursor.fetchone()
    if (improvements):
        return jsonify(improvements[0])
    else:
        context = f"""
        You are an AI expert resume analyzer that always responds in clean nested JSON format. The user will input a resume, in text form, 
        and also a job description, in  text form. You must analyze the resume based on the job descriptions and suggest improvements
        based on the job description, but also general resume improvements.

        Respond in the following format:
        [
            {{
                "suggestion": "What to improve",
                "referenceText": "Exact copy and pasted text from the resume that the suggestion refers to"
            }},
            {{
                "suggestion": "What to improve",
                "referenceText": "Exact copy and pasted text from the resume that the suggestion refers to"
            }},
            ...
        ]
        General Guidelines:
        1. The referenceText must be the EXACT TEXT from the resume, character-for-character. DO not paraphrase or summarize. Do not jump between sections in the resume. 
        2. Only respond with the JSON array and no additional comments, etc.
        """

        cursor.execute("SELECT job_description FROM job_description WHERE session_id = %s", (session_id,))
        job_description = cursor.fetchone()[0]

        cursor.execute("SELECT resume_text FROM resumes WHERE session_id = %s", (session_id,))
        resume = cursor.fetchone()[0]

        prompt = f"""
        -RESUME-
        {resume}


        -JOB DESCIPTION-
        {job_description}
        """

        response = openai(context, prompt)
        json_response = json.loads(response)

        cursor.execute(
        "INSERT INTO resume_improvements (session_id, improvements) VALUES (%s, %s)",
        (session_id, json.dumps(json_response)))
     
        connection.commit()
        cursor.close()
        connection.close()
    
        return jsonify(json_response)

@app.route("/api/extract-text", methods=["POST"])
def extract_text_endpoint():
    
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]

    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    try:
        extracted_text = extract_text(file.stream, file.filename)
        return jsonify({"extracted_text": extracted_text}), 200
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return jsonify({"error": f"An error occurred: {e}"}), 500

if __name__ == "__main__":
    app.run(debug=True, port=8080)